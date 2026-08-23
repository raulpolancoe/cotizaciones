from flask import Flask, request, jsonify, send_file, send_from_directory
import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os
import locale
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import win32com.client as win32
import pythoncom
import unicodedata
import re
from threading import Timer


app = Flask(__name__)

# Rutas
EXCEL_PATH = r"C:/Learning/Cotizaciones/Lista_Precios.xlsx"
PLANTILLA_WORD = r"C:/Learning/Cotizaciones/Archivo Base.docx"

# Añadir bordes a celdas
def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_type in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tcPr.append(border)

# Productos desde Excel
@app.route("/productos")
def productos():
    try:
        df = pd.read_excel(EXCEL_PATH)
        print("📄 Productos cargados desde Excel:")
        print(df)
        productos = df.to_dict(orient="records")
        return jsonify(productos)
    except Exception as e:
        print("❌ ERROR leyendo el Excel:", e)
        return jsonify({"error": str(e)}), 500

# Página principal
@app.route("/")
def index():
    return send_from_directory(".", "cotizaciones.html")

# Generar cotización
@app.route("/generar", methods=["POST"])
def generar():
    try:
        data = request.json
        print("📥 Datos recibidos:", data)
        nombre = data["nombre"]
        items = data["items"]

        # Limpiar nombre para archivo
        def limpiar_nombre(nombre):
            nombre = unicodedata.normalize('NFKD', nombre).encode('ascii', 'ignore').decode('ascii')
            nombre = re.sub(r'[^a-zA-Z0-9_]', '_', nombre)
            return nombre.strip().replace("__", "_").replace(" ", "_")

        safe_name = limpiar_nombre(nombre)
        pdf_out = f"Cotizacion_{safe_name}.pdf"
        docx_out = "CotizacionGenerada.docx"

        # Eliminar archivos antiguos
        for f in [docx_out, pdf_out]:
            if os.path.exists(f):
                try:
                    os.remove(f)
                except Exception as e:
                    print(f"⚠️ No se pudo borrar {f}: {e}")

        # Cargar plantilla
        doc = Document(PLANTILLA_WORD)

        # Fecha en español
        try:
            locale.setlocale(locale.LC_TIME, "es_CO.utf8")
        except locale.Error:
            locale.setlocale(locale.LC_TIME, "es_CO")
        fecha_actual = datetime.today().strftime("%d de %B de %Y").replace(" 0", " ")

        # Reemplazar campos
        for p in doc.paragraphs:
            for run in p.runs:
                if "___ de ___ de ___" in run.text:
                    run.text = run.text.replace("___ de ___ de ___", fecha_actual)
                elif "Señor" in run.text or "Señora" in run.text:
                    run.text = f"Señor(a) {nombre}"
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

        # Llenar tabla
        for table in doc.tables:
            if table.cell(0, 1).text.strip() == "Descripción":
                for _ in range(len(table.rows) - 1):
                    table._tbl.remove(table.rows[1]._tr)
                total = 0
                for item in items:
                    cantidad = int(item['cantidad'])
                    valor = int(item['valor'])
                    subtotal = cantidad * valor
                    row = table.add_row().cells
                    row[0].text = str(cantidad)
                    row[1].text = item['descripcion']
                    row[2].text = f"${valor:,}".replace(",", ".")
                    row[3].text = f"${subtotal:,}".replace(",", ".")
                    total += subtotal
                    for cell in row:
                        set_cell_border(cell)
                        for p in cell.paragraphs:
                            if p.runs:
                                p.runs[0].font.name = 'Calibri'
                                p.runs[0].font.size = Pt(11)
                total_row = table.add_row().cells
                total_row[2].text = "TOTAL"
                total_row[3].text = f"${total:,}".replace(",", ".")
                for cell in total_row:
                    set_cell_border(cell)
                    for p in cell.paragraphs:
                        if p.runs:
                            p.runs[0].font.name = 'Calibri'
                            p.runs[0].font.size = Pt(11)
                break

        # Guardar .docx
        doc.save(docx_out)

        # Convertir a PDF
        try:
            pythoncom.CoInitialize()
            word = win32.Dispatch('Word.Application')
            word.Visible = False
            docx_path = os.path.abspath(docx_out)
            pdf_path = os.path.abspath(pdf_out)

            doc_word = word.Documents.Open(docx_path)
            doc_word.SaveAs(pdf_path, FileFormat=17)
            doc_word.Close()
            word.Quit()
            pythoncom.CoUninitialize()

            # Programar la eliminación de archivos temporales
            def borrar_archivos_temporales():
                for f in [docx_out, pdf_path]:
                    if os.path.exists(f):
                        try:
                            os.remove(f)
                            print(f"🧹 Archivo eliminado: {f}")
                        except Exception as e:
                            print(f"⚠️ No se pudo eliminar {f}: {e}")

            # Programar borrado 5 segundos después
            Timer(5.0, borrar_archivos_temporales).start()

            # Enviar el archivo al cliente
            return send_file(pdf_path, as_attachment=True, download_name=pdf_out, mimetype="application/pdf")


        except Exception as e:
            print("⚠️ Error al convertir con Word:", e)
            return send_file(docx_out, as_attachment=True)

    except Exception as e:
        print("❌ ERROR en generación de cotización:", e)
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(debug=True)
